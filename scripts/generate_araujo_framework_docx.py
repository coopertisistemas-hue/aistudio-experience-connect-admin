from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path("output/doc")
OUTPUT_FILE = OUTPUT_DIR / "ARAUJO-INNOVATION-LAB-FRAMEWORK-ESTRUTURAL-E-OPERACIONAL-v1.docx"


NAVY = "17324D"
TEAL = "2B7A78"
SAND = "F4EFE7"
STONE = "B8B1A6"
LIGHT_TEAL = "E8F4F3"
LIGHT_NAVY = "EEF3F8"
LIGHT_AMBER = "FCF2DC"
AMBER = "C88B2C"
SOFT_GRAY = "F7F7F5"
TEXT = "2F3A45"
MUTED = "687585"


@dataclass
class MatrixRow:
    label: str
    values: list[str]


def set_page_margins(section, top=2.2, bottom=1.9, left=2.3, right=2.0):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D4CB", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, TEAL),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(section, title_text: str):
    header = section.header
    header.is_linked_to_previous = False
    footer = section.footer
    footer.is_linked_to_previous = False

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_run = p.add_run("EEB ARAUJO FIGUEIREDO")
    p_run.font.name = "Aptos"
    p_run.font.size = Pt(8.5)
    p_run.bold = True
    p_run.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("  |  ").font.color.rgb = RGBColor.from_string(STONE)
    r2 = p.add_run(title_text)
    r2.font.name = "Aptos"
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(MUTED)

    f = footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, color in [
        ("Framework Estrutural e Operacional", NAVY),
        ("  |  ", STONE),
        ("Versão 1.0", MUTED),
        ("  |  ", STONE),
        ("Revisão institucional - maio de 2026", MUTED),
        ("  |  Página ", STONE),
    ]:
        run = f.add_run(text)
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(color)
    add_page_number(f)


def add_paragraph(doc: Document, text: str, *, style=None, align=None, bold=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Aptos"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold:
        run.bold = True
    return p


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_TEAL, accent: str = TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, "10")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    r1.font.name = "Aptos"
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor.from_string(accent)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph("")


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = hdr_cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY, "8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        if widths:
            cell.width = Cm(widths[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell, "D7D2C8", "6")
            set_cell_shading(cell, "FFFFFF" if idx else SOFT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = "Aptos"
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor.from_string(TEXT)
            if widths:
                cell.width = Cm(widths[idx])
    doc.add_paragraph("")
    return table


def add_matrix_table(doc: Document, title: str, headers: list[str], rows: list[MatrixRow], caption: str):
    add_paragraph(doc, title, style="Heading 2")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TEAL if idx else NAVY)
        set_cell_border(cell, "D7D2C8", "6")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = "Aptos"
        run.font.size = Pt(8.8)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")

    for row in rows:
        data = [row.label] + row.values
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            cell = cells[idx]
            set_cell_shading(cell, LIGHT_NAVY if idx == 0 else "FFFFFF")
            set_cell_border(cell, "D7D2C8", "6")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.name = "Aptos"
            run.font.size = Pt(8.6)
            run.font.color.rgb = RGBColor.from_string(TEXT)
            if idx == 0:
                run.bold = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_run = cap.add_run(caption)
    cap_run.font.name = "Aptos"
    cap_run.font.italic = True
    cap_run.font.size = Pt(8.5)
    cap_run.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph("")


def add_cover(doc: Document):
    add_paragraph(doc, "LOGOTIPO INSTITUCIONAL", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=TEAL, bold=True)
    add_paragraph(doc, "EEB ARAUJO FIGUEIREDO", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, color=NAVY, bold=True)
    add_paragraph(doc, "Urubici - Santa Catarina", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=MUTED)
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_paragraph(
        doc,
        "FRAMEWORK ESTRUTURAL E OPERACIONAL",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=24,
        color=NAVY,
        bold=True,
    )
    add_paragraph(
        doc,
        "Araujo Innovation Lab",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=18,
        color=TEAL,
        bold=True,
    )
    add_paragraph(
        doc,
        "Documento institucional de organização pedagógica, governança educacional e implementação progressiva.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11.5,
        color=TEXT,
    )
    doc.add_paragraph("")
    add_callout(
        doc,
        "Posicionamento operacional",
        (
            "Este framework consolida a arquitetura institucional do laboratório como modelo executável de inovação "
            "educacional, articulando formação, prática, governança, avaliação e sustentabilidade em linguagem clara, "
            "pedagógica e institucional."
        ),
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_paragraph(
        doc,
        "Assinatura institucional",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9.5,
        color=MUTED,
        bold=True,
    )
    add_paragraph(
        doc,
        "EEB Araujo Figueiredo  |  Laboratório de Inovação Educacional  |  Maio de 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=9,
        color=MUTED,
    )


def add_index(doc: Document):
    doc.add_page_break()
    add_paragraph(doc, "Índice Executivo", style="Heading 1")
    add_paragraph(
        doc,
        "Estrutura de leitura pensada para circulação institucional, apresentação executiva e consulta pedagógica.",
        color=MUTED,
    )
    rows = [
        ["01", "Identidade e propósito do Araujo Innovation Lab", "Visão institucional, objetivos e posicionamento."],
        ["02", "Arquitetura estrutural do laboratório", "Eixos de funcionamento, frentes de atuação e organização prática."],
        ["03", "Matriz 6D e progressão formativa", "Jornada pedagógica integrada e lógica de desenvolvimento."],
        ["04", "Matriz interdisciplinar", "Conexão entre áreas do currículo e práticas do laboratório."],
        ["05", "Governança educacional", "Coordenação, fluxos de comunicação, continuidade e sustentabilidade."],
        ["06", "Laboratório de Marketing Digital", "Áreas práticas, entregas orientadas e repertório de mercado."],
        ["07", "Fluxo de experiências e conexão com o território", "Visitas, observações, projetos e integração empresarial."],
        ["08", "Roadmap de implementação progressiva", "MVP, consolidação, expansão e escalabilidade responsável."],
        ["09", "Sistema de avaliação e feedback", "Acompanhamento contínuo, conselho de classe e evolução dos estudantes."],
        ["10", "Diretrizes de fechamento e exportação institucional", "Leitura final, circulação executiva e PDF."],
    ]
    add_simple_table(doc, ["Seção", "Bloco", "Foco de leitura"], rows, widths=[2.0, 7.4, 7.2])


def add_section_intro(doc: Document):
    add_paragraph(doc, "1. Identidade e Propósito do Araujo Innovation Lab", style="Heading 1")
    add_paragraph(
        doc,
        "O Araujo Innovation Lab nasce como estrutura educacional aplicada para conectar formação escolar, repertório "
        "contemporâneo, cultura digital e desenvolvimento territorial. Seu papel não é apenas ampliar experiências; "
        "é organizar um ambiente permanente de aprendizagem orientada por prática, investigação, autoria e visão de futuro."
    )
    add_paragraph(
        doc,
        "Em vez de operar como atividade complementar isolada, o laboratório se posiciona como eixo institucional de "
        "integração pedagógica. Ele aproxima escola, estudantes, professores e ecossistema local por meio de uma dinâmica "
        "clara: aprender, experimentar, produzir, apresentar, revisar e evoluir."
    )
    add_callout(
        doc,
        "Princípio institucional",
        "O laboratório deve comunicar viabilidade real. Cada ação, espaço, parceria e entrega precisa ser compreensível, replicável e sustentável dentro da rotina escolar.",
    )


def add_structural_architecture(doc: Document):
    add_paragraph(doc, "2. Arquitetura Estrutural do Laboratório", style="Heading 1")
    add_paragraph(
        doc,
        "A estrutura operacional do laboratório foi desenhada para manter clareza de função, flexibilidade pedagógica "
        "e continuidade institucional. O objetivo é permitir leitura imediata dos eixos de trabalho sem criar sensação de "
        "manual técnico ou excesso de complexidade."
    )
    rows = [
        ["Núcleo formativo", "Base conceitual, cultura digital, repertório, comportamento ético e visão de projeto."],
        ["Núcleo de prática orientada", "Execução supervisionada, experimentação, prototipagem e entregas guiadas."],
        ["Núcleo de integração curricular", "Articulação com componentes acadêmicos e projetos interdisciplinares."],
        ["Núcleo de conexão externa", "Território, empresas locais, observação profissional e experiências de mercado."],
        ["Núcleo de governança e continuidade", "Planejamento, acompanhamento, indicadores, revisão e sustentabilidade."],
    ]
    add_simple_table(doc, ["Eixo estrutural", "Função operacional"], rows, widths=[5.2, 11.4])
    add_callout(
        doc,
        "Leitura executiva",
        "A arquitetura do laboratório funciona como um sistema em camadas: forma, pratica, integra, conecta e sustenta.",
        fill=LIGHT_NAVY,
        accent=NAVY,
    )


def add_6d_matrix(doc: Document):
    add_paragraph(doc, "3. Matriz 6D e Progressão Formativa", style="Heading 1")
    add_paragraph(
        doc,
        "A Matriz 6D organiza a experiência do estudante em uma lógica progressiva e inteligível. O foco não está em "
        "complexificar a leitura, mas em mostrar como a evolução pedagógica acontece de forma prática, cumulativa e "
        "institucionalmente acompanhada."
    )
    headers = ["Dimensão", "Sentido pedagógico", "Aplicação no laboratório", "Indicadores de progresso"]
    rows = [
        MatrixRow("Descobrir", ["Contato inicial com repertório e possibilidades.", "Imersões, referências, observações e leitura de contexto.", "Curiosidade, participação, registro de interesse."]),
        MatrixRow("Diagnosticar", ["Leitura crítica de problemas e oportunidades.", "Mapeamento de desafios, análise de casos e escuta orientada.", "Capacidade de identificar necessidades e formular perguntas."]),
        MatrixRow("Desenhar", ["Estruturação de ideias e propostas.", "Planejamento de ações, definição de solução e roteiros de execução.", "Organização, coerência e intencionalidade das propostas."]),
        MatrixRow("Desenvolver", ["Construção prática de entregas.", "Produção de peças, testes, experimentos e protótipos.", "Evolução técnica, autonomia progressiva e consistência."]),
        MatrixRow("Demonstrar", ["Apresentação, socialização e argumentação.", "Mostras, pitches, relatórios visuais e portfólios.", "Clareza de comunicação, segurança e autoria."]),
        MatrixRow("Direcionar", ["Revisão, continuidade e visão futura.", "Feedback, replanejamento, consolidação e novos ciclos.", "Capacidade de refletir, melhorar e sustentar a trajetória."]),
    ]
    add_matrix_table(
        doc,
        "Matriz 6D",
        headers,
        rows,
        "Figura operacional 1 - Progressão formativa em linguagem institucional, com foco em clareza e aplicabilidade pedagógica.",
    )


def add_interdisciplinary_matrix(doc: Document):
    add_paragraph(doc, "4. Matriz Interdisciplinar", style="Heading 1")
    add_paragraph(
        doc,
        "A interdisciplinaridade aparece aqui como articulação natural entre áreas do conhecimento, e não como sobrecarga "
        "curricular. A matriz abaixo ajuda a visualizar como cada componente contribui para experiências concretas do laboratório."
    )
    headers = ["Área", "Contribuição central", "Expressão no laboratório"]
    rows = [
        ["Português", "Narrativa, argumentação, leitura crítica e produção textual.", "Copywriting, roteiros, apresentações, relatórios e portfólio."],
        ["Matemática", "Lógica, mensuração, interpretação de dados e projeção.", "Analytics, métricas, custos, desempenho e leitura de indicadores."],
        ["Artes", "Sensibilidade estética, composição visual e linguagem criativa.", "Branding, identidade visual, audiovisual e direção de peças."],
        ["Tecnologia", "Ferramentas digitais, experimentação e fluência operacional.", "IA aplicada, automações simples, produção multimídia e plataformas."],
        ["Sociologia", "Leitura de contexto, território, comportamento e cultura.", "Mapeamento de público, observação social e sentido comunitário."],
        ["Administração", "Organização, planejamento, processos e visão de sustentabilidade.", "Projetos supervisionados, cronogramas, priorização e operação."],
        ["Comunicação", "Posicionamento, clareza de mensagem e interação.", "Campanhas, presença digital, reputação e comunicação institucional."],
        ["Projeto de Vida", "Autoconhecimento, protagonismo e visão de trajetória.", "Portfólio, escolhas de interesse, metas e conexão com futuro profissional."],
    ]
    add_simple_table(doc, headers, rows, widths=[3.2, 6.0, 7.4])
    add_callout(
        doc,
        "Leitura pedagógica",
        "A matriz interdisciplinar deve apoiar o planejamento dos professores e dar previsibilidade às famílias sobre o valor educacional do laboratório.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )


def add_governance(doc: Document):
    add_paragraph(doc, "5. Governança Educacional", style="Heading 1")
    add_paragraph(
        doc,
        "A governança do Araujo Innovation Lab precisa transmitir coordenação, continuidade e responsabilidade institucional. "
        "O modelo abaixo privilegia simplicidade visual, papéis claros e comunicação sustentável ao longo do ano letivo."
    )
    rows = [
        ["Direção escolar", "Patrocínio institucional, legitimidade e alinhamento estratégico."],
        ["Coordenação pedagógica", "Integração curricular, acompanhamento docente e continuidade pedagógica."],
        ["Coordenação do laboratório", "Organização operacional, agenda, articulações externas e curadoria de experiências."],
        ["Professores referência", "Mediação de projetos, acompanhamento formativo e integração com áreas do conhecimento."],
        ["Estudantes", "Participação ativa, produção supervisionada, autoria e portfólio."],
        ["Parceiros e território", "Contexto real, observação profissional, repertório e conexões de futuro."],
    ]
    add_simple_table(doc, ["Nível de governança", "Responsabilidade principal"], rows, widths=[5.0, 11.6])
    add_paragraph(doc, "Fluxo de comunicação e continuidade", style="Heading 2")
    flow_rows = [
        ["Planejamento", "Direção + coordenação pedagógica + coordenação do laboratório definem prioridades do ciclo."],
        ["Execução", "Professores e estudantes desenvolvem experiências com acompanhamento claro."],
        ["Registro", "Entregas, evidências, observações e aprendizados são consolidados de forma organizada."],
        ["Revisão", "Os resultados retornam ao conselho pedagógico e à coordenação para ajuste de rota."],
        ["Continuidade", "As decisões do ciclo alimentam o planejamento seguinte e garantem sustentabilidade."],
    ]
    add_simple_table(doc, ["Etapa", "Lógica operacional"], flow_rows, widths=[3.2, 13.4])


def add_lab_structure(doc: Document):
    add_paragraph(doc, "6. Laboratório de Marketing Digital", style="Heading 1")
    add_paragraph(
        doc,
        "O Laboratório de Marketing Digital constitui uma frente prática de alta legibilidade para estudantes, professores e "
        "parceiros. A organização abaixo evita dispersão e comunica um conjunto coerente de áreas, cada uma com função clara "
        "dentro da experiência formativa."
    )
    headers = ["Área prática", "Aplicação pedagógica", "Entregas possíveis"]
    rows = [
        ["Branding", "Compreensão de identidade, posicionamento e consistência visual.", "Naming, identidade, peças institucionais e guias simples."],
        ["IA aplicada", "Uso orientado de ferramentas para pesquisa, ideação e produtividade.", "Prompts supervisionados, apoio criativo e síntese assistida."],
        ["Google Meu Negócio", "Presença local, reputação e visibilidade territorial.", "Mapeamento, perfis simulados, diagnósticos e boas práticas."],
        ["Copywriting", "Escrita persuasiva, clareza de mensagem e estrutura narrativa.", "Legendas, chamadas, textos institucionais e roteiros curtos."],
        ["Audiovisual", "Planejamento visual, captação e edição básica orientada.", "Vídeos curtos, bastidores, depoimentos e registros de projeto."],
        ["Analytics", "Leitura de dados, métricas e interpretação de desempenho.", "Painéis simples, resumos executivos e aprendizagem baseada em evidências."],
        ["Reputação digital", "Comportamento de marca, escuta e percepção pública.", "Protocolos de resposta, análise de comentários e observação crítica."],
    ]
    add_simple_table(doc, headers, rows, widths=[4.0, 6.0, 6.6])


def add_experience_market_flow(doc: Document):
    add_paragraph(doc, "7. Fluxo de Experiências e Conexão com o Mercado", style="Heading 1")
    add_paragraph(
        doc,
        "A aproximação com o mundo do trabalho deve acontecer com progressão, segurança pedagógica e leitura realista de "
        "possibilidades. O fluxo a seguir organiza experiências de forma crescente, evitando promessas excessivas e reforçando "
        "a maturidade operacional do programa."
    )
    rows = [
        ["01. Sensibilização", "Introduções guiadas ao território, repertório de setores e observação de referências."],
        ["02. Visitas técnicas", "Contato presencial com ambientes organizacionais e leitura estruturada de processos."],
        ["03. Observação empresarial", "Análise supervisionada de práticas reais, comunicação e rotina de operação."],
        ["04. Projetos supervisionados", "Desafios orientados com escopo pedagógico claro e devolutiva de professores."],
        ["05. Experiências práticas", "Produções autorais e aplicações contextualizadas com segurança institucional."],
        ["06. Portfólio e projeção", "Consolidação de trajetórias, evidências de aprendizagem e visão de continuidade."],
    ]
    add_simple_table(doc, ["Etapa de experiência", "Sentido formativo"], rows, widths=[5.0, 11.6])
    add_callout(
        doc,
        "Conexão com o território",
        "A integração com negócios locais deve fortalecer aprendizagem e pertencimento, sem transformar o laboratório em prestação de serviço improvisada.",
    )


def add_roadmap(doc: Document):
    add_paragraph(doc, "8. Implementação Progressiva", style="Heading 1")
    add_paragraph(
        doc,
        "O roadmap foi estruturado para comunicar avanço realista. A implementação progressiva protege a escola de sobrecarga, "
        "preserva qualidade pedagógica e permite amadurecimento institucional a cada fase."
    )
    rows = [
        ["Fase 1 - MVP operacional", "Organizar base, pilotar dinâmica, validar rotina e consolidar linguagem institucional.", "Escopo controlado, poucas frentes, alta observação e rápida aprendizagem."],
        ["Fase 2 - Consolidação", "Ajustar governança, ampliar integração curricular e estabilizar calendário do laboratório.", "Mais previsibilidade, evidências consistentes e melhoria de processos."],
        ["Fase 3 - Expansão qualificada", "Ampliar projetos, parceiros e repertório com segurança pedagógica.", "Crescimento gradual sem perda de clareza nem excesso de complexidade."],
        ["Fase 4 - Sustentabilidade e legado", "Formalizar continuidade, memória institucional e transferência de aprendizado.", "Modelo replicável, estável e reconhecido pela comunidade escolar."],
    ]
    add_simple_table(doc, ["Fase", "Foco principal", "Critério de maturidade"], rows, widths=[4.4, 6.0, 6.2])


def add_evaluation(doc: Document):
    add_paragraph(doc, "9. Sistema de Avaliação e Feedback", style="Heading 1")
    add_paragraph(
        doc,
        "A avaliação no laboratório deve reforçar processo, autoria e evolução, evitando tanto a superficialidade quanto a "
        "hipertrofia burocrática. O acompanhamento precisa ser claro para estudantes, professores e conselho de classe."
    )
    rows = [
        ["Observação inicial", "Mapear repertório, interesses e pontos de apoio pedagógico."],
        ["Acompanhamento de percurso", "Registrar participação, consistência, colaboração e aplicação prática."],
        ["Feedback formativo", "Oferecer devolutivas curtas, orientadas e acionáveis ao longo do ciclo."],
        ["Revisão com conselho de classe", "Conectar evidências do laboratório ao desenvolvimento acadêmico mais amplo."],
        ["Portfólio e síntese", "Consolidar produções, reflexões e marcos de evolução do estudante."],
    ]
    add_simple_table(doc, ["Componente", "Função pedagógica"], rows, widths=[5.0, 11.6])
    add_paragraph(doc, "Ciclos de feedback", style="Heading 2")
    feedback_rows = [
        ["Semanal", "Checagens breves de andamento, rotina e participação."],
        ["Mensal", "Leitura de entregas, qualidade de processo e redirecionamentos necessários."],
        ["Bimestral", "Síntese pedagógica, evidências principais e comunicação ampliada."],
        ["Semestral", "Consolidação institucional, avaliação de impacto e próximos ajustes."],
    ]
    add_simple_table(doc, ["Ritmo", "Objetivo"], feedback_rows, widths=[3.2, 13.4])


def add_executive_blocks(doc: Document):
    add_paragraph(doc, "10. Blocos Executivos de Leitura", style="Heading 1")
    add_paragraph(
        doc,
        "Ao longo do framework, determinados blocos devem cumprir função editorial específica. Eles ajudam a manter "
        "respiração visual, clareza executiva e legibilidade pedagógica."
    )
    add_callout(
        doc,
        "Observação pedagógica",
        "Usar para explicitar sentido formativo, cuidado com o estudante e valor educacional das decisões operacionais.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    add_callout(
        doc,
        "Princípio operacional",
        "Usar para sintetizar critérios de execução, limites de escopo e decisões que sustentam o funcionamento do laboratório.",
        fill=LIGHT_NAVY,
        accent=NAVY,
    )
    add_callout(
        doc,
        "Nota de implementação",
        "Usar para registrar pontos de atenção, cadência de implantação e dependências práticas do ciclo seguinte.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )
    add_callout(
        doc,
        "Destaque institucional",
        "Usar para reforçar visão, legitimidade, continuidade e alinhamento com a missão da escola.",
        fill=SOFT_GRAY,
        accent=NAVY,
    )


def add_pdf_export(doc: Document):
    add_paragraph(doc, "11. Preparação para Exportação em PDF e Circulação Institucional", style="Heading 1")
    add_paragraph(
        doc,
        "A versão em Google Docs deve preservar leitura estável quando exportada em PDF para circulação com direção, equipe "
        "pedagógica, parceiros e comunidade institucional. A preparação final precisa garantir consistência visual, paginação "
        "controlada e integridade dos quadros e matrizes."
    )
    rows = [
        ["Margens e respiros", "Preservar áreas de leitura amplas e evitar elementos colados às bordas."],
        ["Paginação", "Conferir sequência estável e manter páginas de abertura com leitura limpa."],
        ["Tabelas e matrizes", "Evitar quebras confusas e validar legibilidade de cada quadro no PDF final."],
        ["Cabeçalhos e rodapés", "Manter assinatura institucional discreta, consistente e não invasiva."],
        ["Exportação", "Gerar PDF com qualidade padrão, sem compressão agressiva e com revisão final de leitura."],
    ]
    add_simple_table(doc, ["Item de verificação", "Diretriz de fechamento"], rows, widths=[5.2, 11.4])
    add_callout(
        doc,
        "Polimento final",
        "O documento deve sair para apresentação institucional com sensação de maturidade, clareza e viabilidade real - nunca como rascunho técnico ou relatório de consultoria.",
        fill=LIGHT_AMBER,
        accent=AMBER,
    )


def add_closing(doc: Document):
    add_paragraph(doc, "12. Fechamento Executivo", style="Heading 1")
    add_paragraph(
        doc,
        "O Araujo Innovation Lab se apresenta, neste framework, como uma proposta de transformação educacional ancorada em "
        "governança, progressão pedagógica, experimentação responsável e continuidade institucional. Seu valor está menos na "
        'promessa de inovação abstrata e mais na capacidade de organizar um ecossistema escolar que aprende, produz, reflete e evolui com consistência.'
    )
    add_paragraph(
        doc,
        "Ao comunicar estrutura, ritmo, práticas, avaliação e implementação progressiva, este documento estabelece um padrão "
        "de leitura que combina sofisticação executiva e clareza operacional. O resultado esperado é um framework que pode "
        "ser apresentado, discutido, ajustado e colocado em prática com legitimidade institucional."
    )
    add_callout(
        doc,
        "Síntese institucional",
        "Mais do que um programa complementar, o laboratório é um eixo de cultura educacional aplicada, orientado por maturidade pedagógica, inovação responsável e sustentabilidade de longo prazo.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )


def create_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_base_styles(doc)
    set_page_margins(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True

    add_cover(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(section)
    add_header_footer(section, "Araujo Innovation Lab")

    add_index(doc)
    add_section_intro(doc)
    add_structural_architecture(doc)
    add_6d_matrix(doc)
    add_interdisciplinary_matrix(doc)
    add_governance(doc)
    add_lab_structure(doc)
    add_experience_market_flow(doc)
    add_roadmap(doc)
    add_evaluation(doc)
    add_executive_blocks(doc)
    add_pdf_export(doc)
    add_closing(doc)

    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    path = create_document()
    print(path)
